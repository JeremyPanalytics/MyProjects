import docx
import pandas as pd
import re

# file_name = input("what is the file name and the extension?")
file_name = "test_file.docx"
doc = docx.Document(file_name)


# number of paragraphs
all_paras = doc.paragraphs
print(f"number of paragraphs: {len(all_paras)}")

# convert to list
paragraphs_all = [para.text for para in doc.paragraphs if para.text != ""]
# print(paragraphs_all)
print(f"number of paragraphs without empty spaces: {len(paragraphs_all)}")

# conversion to table
paragraph_IDs = range(0, len(paragraphs_all))

df = pd.DataFrame(paragraphs_all, columns=["Paragraph"], index=range(0,len(paragraphs_all)))
# print(df)

# locate main text to run search
# starts from the second Introduction
# ends before the last Introduction in the provisional data plan

search_term = "Introduction"
search_list = []


def checker(x):
    if x == None:
        return 0
    else:
        return 1


for i in df.Paragraph:
    a = re.search("Introduction", i)
    search_list.append(checker(a))

search_list = pd.Series(search_list, name="Intro_Loc")
# print(search_list)

df = pd.merge(df, search_list, left_index=True, right_index=True)

# to start from the first introduction, I will need to set an index
# text_has_content = input("Does text have a content page? Yes or No: ").lower()
text_has_content = "yes" #default

if text_has_content == "yes":
    presence_of_contentpage = 1
elif text_has_content == "no":
    presence_of_contentpage = 0

starting_index = df[df.Intro_Loc == 1].index[presence_of_contentpage]

df_working = df[starting_index:].drop(columns="Intro_Loc", axis=1)
# print(df_working.head())

#df_working now starts from introduction

# TODO 2: Search for text between ( and )
print(f"Current Number of working paragraphs in df_working: {len(df_working)}")
all_parentheses = []
for i in df_working.Paragraph:
    a = re.findall('\(.*?\)', i)
    all_parentheses.append(a)

# note: [] or blanks indicate that paragraph has no parentheses
# remove empty lists
all_parentheses = [list for list in all_parentheses if list !=[]]

# combine all parentheses to one list:
all_parentheses = [item for sublist in all_parentheses for item in sublist]

# TODO 3: Remove parentheses with text only

pattern = r'\(([,a-zA-Z, -.""\[|\]/\+]+)\)'
text_only = []
for i in all_parentheses:
     a = re.findall(pattern, i)
     text_only.append(a)

citations = pd.DataFrame(all_parentheses, columns=["citations"])
text_only = pd.DataFrame(text_only, columns=["text_only"])

# this is so that I could anti-join the columns
test_df = pd.merge(citations, text_only,
                   how="outer",
                   left_index=True,
                   right_index=True,
                   )

# print(text_only)
final = test_df[test_df.text_only.isnull()].citations
final_list = []
for i in final:
    a = i.strip("()")
    final_list.append(a)
final_list = pd.DataFrame(final_list, columns=["check here"])


# TODO 4: I think the best way forward is to search for (text... \4d)?
# i updated this about 2 weeks after the last section. this works better!

# essentially, this search finds strings that end with years between 1000 - 2999 and the end of the parentheses
pattern = r'.*[1-2][09][0-9][0-9]\)$'
final_list2 = []
for i in all_parentheses:
    b = re.findall(pattern, i)
    for j in b:
        c = j.strip("()")
        final_list2.append(c)

final_list2 = [list for list in final_list2 if list != []]
citations_only = pd.DataFrame(final_list2, columns=["citations"])

# run to see output
#citations_only.to_csv("final_test.csv", index=False, encoding="ANSI")
